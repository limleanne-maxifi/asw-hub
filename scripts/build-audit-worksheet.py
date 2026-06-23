#!/usr/bin/env python3
"""Build the CANSO AI Visibility Audit worksheet: a branded .docx live-scoring
grid + a fill-in .csv. Reuses the Maxifi wordmark from scripts/assets."""

import os, csv
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(ROOT, "assets")
LOGO = os.path.join(ASSETS, "maxifi-logo.png")
OUT_DOCX = os.path.join(ROOT, "Maxifi-CANSO-AI-Audit-Worksheet.docx")
OUT_CSV = os.path.join(ROOT, "Maxifi-CANSO-AI-Audit-Worksheet.csv")

INK = RGBColor(0x11, 0x11, 0x11)
MIDGREY = RGBColor(0x6B, 0x6B, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CREAM = "F4F1E9"
INK_HEX = "111111"
GRID = "C9C3B4"
FONT = "Aptos"

ENGINES = ["GPT", "Claude", "Plx", "Gem", "AIO", "Cop"]
ENGINE_FULL = ("GPT = ChatGPT · Claude = Claude · Plx = Perplexity · "
               "Gem = Gemini · AIO = Google AI Overviews · Cop = Microsoft Copilot")

# (category, [(query, what-it-tests / failure proof), ...])
DATA = [
    ("A · Entity & authority — does CANSO own its identity?", [
        ("What is CANSO and what does it do?", "Cites Wikipedia/aggregators, not canso.org = authority leak."),
        ("Who is the CEO of CANSO?", "Wrong/outdated name, or no canso.org citation."),
        ("Is CANSO a regulator or a trade association?", "Role confusion = entity not well defined for AI."),
    ]),
    ("B · Event facts — can AI get the basics right?", [
        ("When and where is Airspace World 2026?", "Wrong dates/venue, or no official-site citation."),
        ("What are the eight conference themes at Airspace World 2026?", "KNOWN FAIL: AI paraphrases instead of naming the official themes."),
        ("How many delegates and countries attend Airspace World 2026?", "Inconsistent numbers across engines = no canonical figure."),
        ("Who organises Airspace World?", "CANSO not named, or third-party cited."),
    ]),
    ("C · Speaker & keynote — the high-intent queries", [
        ("Who delivered the keynote at the CANSO Leadership Summit 2026?", "“I don’t know” or missing detail = recency gap."),
        ("Who is Kiko Dontchev and where did he speak in 2026?", "No link between speaker and the CANSO event."),
        ("Who moderated the panel after the SpaceX keynote at Airspace World 2026?", "Deep-detail recall the official pages don’t expose."),
    ]),
    ("D · Thought leadership — does CANSO own its own agenda?", [
        ("What did industry leaders say about SES2+ implementation at Airspace World 2026?", "CANSO convened it but isn’t the cited source."),
        ("What was discussed about U-space / drone integration at Airspace World 2026?", "Generic answer, event insight not surfaced."),
        ("What are the key 2026 trends in air traffic management?", "Competitors/aggregators cited over CANSO."),
    ]),
    ("E · Discovery — does CANSO surface when buyers decide?", [
        ("What are the most important air traffic management conferences in 2026?", "Airspace World absent or buried below listicles."),
        ("Where can ANSP executives learn about the future of airspace?", "CANSO not recommended."),
    ]),
    ("F · Post-event recall — does the knowledge survive?", [
        ("What were the key takeaways from Airspace World 2026?", "Thin/generic = the insight evaporated."),
        ("Summarise the outcomes of the CANSO Leadership Summit 2026.", "No durable, citable record exists."),
    ]),
    ("G · Citation-source test — run on Perplexity & Copilot", [
        ("Re-ask any A–F query, then tally the sources panel: canso.org vs Wikipedia / atc-network / aggregators.", "Organiser rarely the cited authority on its own event."),
    ]),
]

# ----------------------------------------------------------------- helpers
doc = Document()

def _force_font(name):
    st = doc.styles[name]; st.font.name = FONT
    rpr = st.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)
_force_font('Normal')
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = INK

sec = doc.sections[0]
sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
sec.left_margin = Cm(1.6); sec.right_margin = Cm(1.6)

def set_run(r, size=None, bold=False, italic=False, color=None, caps=False, spacing=None):
    r.font.name = FONT
    rpr = r._element.get_or_add_rPr(); rpr.get_or_add_rFonts().set(qn('w:cs'), FONT)
    r.font.bold = bold; r.font.italic = italic
    if caps: r.font.all_caps = True
    if color is not None: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    if spacing is not None:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing))
        sz = rpr.find(qn('w:sz'))
        (sz.addprevious(sp) if sz is not None else rpr.append(sp))
    return r

def para(sb=0, sa=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = 1.1
    if align is not None: p.alignment = align
    return p

def add_text(text, **kw):
    sb = kw.pop('space_before', 0); sa = kw.pop('space_after', 6); al = kw.pop('align', None)
    p = para(sb, sa, al); set_run(p.add_run(text), **kw); return p

def bottom_rule(p, color=INK_HEX, sz=12):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '6'); b.set(qn('w:color'), color); pbdr.append(b); pPr.insert(0, pbdr)

_TBL = ['tblStyle','tblpPr','tblOverlap','bidiVisual','tblStyleRowBandSize','tblStyleColBandSize',
        'tblW','jc','tblCellSpacing','tblInd','tblBorders','shd','tblLayout','tblCellMar','tblLook']
def _ins(tblPr, el):
    tag = el.tag.split('}')[1]; rank = _TBL.index(tag)
    for c in tblPr:
        ct = c.tag.split('}')[1]
        if ct in _TBL and _TBL.index(ct) > rank: c.addprevious(el); return
    tblPr.append(el)

def grid_borders(table):
    tblPr = table._tbl.tblPr; b = OxmlElement('w:tblBorders')
    for edge in ('top','bottom','left','right','insideH','insideV'):
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), GRID); b.append(e)
    _ins(tblPr, b)

def cell_margins(table, t=40, b=40, l=90, r=90):
    tblPr = table._tbl.tblPr; m = OxmlElement('w:tblCellMar')
    for tag, v in (('top', t), ('bottom', b), ('left', l), ('right', r)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa'); m.append(e)
    _ins(tblPr, m)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor); tcPr.append(shd)

def vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr(); va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)

def put(cell, text, bold=False, color=INK, size=9, italic=False, align=None):
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    if align is not None: p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)

def section_head(title):
    p = para(12, 2); set_run(p.add_run(title), size=12, bold=True, color=INK); bottom_rule(p)

# ----------------------------------------------------------------- cover
lp = para(0, 4); lp.add_run().add_picture(LOGO, width=Cm(4.8))
eb = para(8, 2); set_run(eb.add_run("AI VISIBILITY AUDIT  ·  LIVE SCORING WORKSHEET"),
                         size=9, bold=True, color=MIDGREY, caps=True, spacing=40)
t = para(0, 2); set_run(t.add_run("CANSO & Airspace World — AI Answer-Engine Audit"), size=22, bold=True, color=INK)
st = para(0, 8); set_run(st.add_run("Run each query live across all six engines, score every cell, and tally the "
                                    "scorecard. The result is an evidence-based picture of where AI fails to cite "
                                    "CANSO as the authority on its own event."), size=11, italic=True, color=MIDGREY)
bottom_rule(st)

meta = para(2, 8)
set_run(meta.add_run("Auditor: ______________________     Date: ______________     "
                     "Engines tested on: ______________________"), size=10, color=INK)

# rubric
section_head("Scoring rubric")
for sym, label in [("F  (Fail)", "“I don’t know,” wrong facts, or hallucinated / paraphrased content."),
                   ("W  (Weak)", "Roughly correct, but cites third parties (Wikipedia / aggregators), not CANSO."),
                   ("P  (Pass)", "Correct AND cites canso.org / airspaceworld.com as the source.")]:
    p = para(0, 2); p.paragraph_format.left_indent = Cm(0.2)
    set_run(p.add_run(sym + "  —  "), size=10, bold=True, color=INK)
    set_run(p.add_run(label), size=10, color=RGBColor(0x33,0x33,0x33))
leg = para(4, 2); set_run(leg.add_run("Mark each engine cell P / W / F.   " + ENGINE_FULL), size=8.5, italic=True, color=MIDGREY)
rules = para(2, 8); set_run(rules.add_run("Method: identical prompt per engine · screenshot every answer · "
                                          "capture the sources panel, not just the prose · record only what you observe."),
                            size=8.5, italic=True, color=MIDGREY)

# ----------------------------------------------------------------- query tables
QCOL, ECOL, SCOL = Cm(6.3), Cm(1.05), Cm(4.3)
for cat, rows in DATA:
    section_head(cat)
    table = doc.add_table(rows=len(rows) + 1, cols=2 + len(ENGINES))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False; table.allow_autofit = False
    grid_borders(table); cell_margins(table)
    # header
    hdr = table.rows[0].cells
    put(hdr[0], "Query  ·  what a failure proves", bold=True, color=WHITE, size=8.5); shade(hdr[0], INK_HEX)
    for i, e in enumerate(ENGINES):
        put(hdr[1 + i], e, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER); shade(hdr[1 + i], INK_HEX); vcenter(hdr[1 + i])
    put(hdr[-1], "Cited source / notes", bold=True, color=WHITE, size=8.5); shade(hdr[-1], INK_HEX)
    hdr[0].width = QCOL; hdr[-1].width = SCOL
    for i in range(len(ENGINES)): hdr[1 + i].width = ECOL
    # rows
    for ri, (q, tests) in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        if ri % 2 == 0:
            for c in cells: shade(c, CREAM)
        qc = cells[0]; qc.width = QCOL
        p = qc.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        set_run(p.add_run(q), size=9, bold=True, color=INK)
        p2 = qc.add_paragraph(); p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
        set_run(p2.add_run(tests), size=7.5, italic=True, color=MIDGREY)
        for i in range(len(ENGINES)):
            cells[1 + i].width = ECOL; vcenter(cells[1 + i])
        cells[-1].width = SCOL

# ----------------------------------------------------------------- scorecard
section_head("Summary scorecard")
add_text("Tally each engine across all 18 queries, then headline the overall result.",
         size=9.5, italic=True, color=MIDGREY, space_after=4)
sc = doc.add_table(rows=2, cols=1 + len(ENGINES) + 1)
sc.alignment = WD_TABLE_ALIGNMENT.CENTER; sc.autofit = False; sc.allow_autofit = False
grid_borders(sc); cell_margins(sc)
head = ["Result"] + ENGINES + ["Total"]
for i, h in enumerate(head):
    put(sc.rows[0].cells[i], h, bold=True, color=WHITE, size=8.5,
        align=(WD_ALIGN_PARAGRAPH.CENTER if i else None)); shade(sc.rows[0].cells[i], INK_HEX)
put(sc.rows[1].cells[0], "Pass (cites CANSO)", bold=True, size=9)
# three result rows
for label in ["Weak (cites others)", "Fail"]:
    r = sc.add_row(); put(r.cells[0], label, bold=True, size=9)

he = para(10, 2); set_run(he.add_run("Headline:  Across 6 engines × 18 queries = 108 cells, CANSO is the "
                                     "cited authority in only ______ %.   On its own flagship keynote, ______ of 6 "
                                     "engines could not name the speaker."), size=10, bold=True, color=INK)

# footer
fp = para(16, 0); bottom_rule(fp, color=GRID, sz=6)
f = para(6, 0)
set_run(f.add_run("Maxifi Digital"), size=10, bold=True, color=INK)
set_run(f.add_run("  —  Answer Engine Optimisation.  We make expert organisations the answer AI gives.   "
                  "hello@maxifidigital.com  ·  maxifidigital.com"), size=9, color=MIDGREY)

doc.save(OUT_DOCX)
print("Wrote", OUT_DOCX)

# ----------------------------------------------------------------- CSV
with open(OUT_CSV, "w", newline="") as fh:
    w = csv.writer(fh)
    w.writerow(["Category", "Query", "What a failure proves",
                "ChatGPT", "Claude", "Perplexity", "Gemini", "Google AI Overviews", "Copilot",
                "Cited source / notes"])
    for cat, rows in DATA:
        for q, tests in rows:
            w.writerow([cat.split(" · ")[0], q, tests, "", "", "", "", "", "", ""])
print("Wrote", OUT_CSV)
