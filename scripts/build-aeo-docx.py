#!/usr/bin/env python3
"""Build the Maxifi Digital-branded Event & Conference AEO Checklist as a .docx.
Recreates the MAXIFI DIGITAL wordmark (original attachment binary not available
on disk) and embeds it. Swap assets/maxifi-logo.png for the official file and
re-run to rebuild with the real logo."""

import os
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(ROOT, "assets")
os.makedirs(ASSETS, exist_ok=True)
LOGO = os.path.join(ASSETS, "maxifi-logo.png")
OUT = os.path.join(ROOT, "Maxifi-AEO-Checklist.docx")

INK = RGBColor(0x11, 0x11, 0x11)
CREAM = "F4F1E9"
INK_HEX = "111111"
MIDGREY = RGBColor(0x6B, 0x6B, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Aptos"

# ---------------------------------------------------------------- logo
def build_logo():
    bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    mono = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf"
    W, H = 2400, 760
    img = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    black = (17, 17, 17, 255)

    def draw_tracked(text, font, y, track, target_center=W // 2):
        widths = [d.textbbox((0, 0), ch, font=font)[2] for ch in text]
        total = sum(widths) + track * (len(text) - 1)
        x = target_center - total // 2
        for ch, w in zip(text, widths):
            d.text((x, y), ch, font=font, fill=black)
            x += w + track

    f_max = ImageFont.truetype(bold, 360)
    f_dig = ImageFont.truetype(mono, 150)
    draw_tracked("MAXIFI", f_max, 40, 24)
    draw_tracked("DIGITAL", f_dig, 470, 96)
    img.save(LOGO)

build_logo()

# ---------------------------------------------------------------- helpers
doc = Document()

# default font = Aptos everywhere
def _force_font(style_name):
    st = doc.styles[style_name]
    st.font.name = FONT
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), FONT)

for s in ('Normal',):
    _force_font(s)
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = INK

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

def set_run(r, size=None, bold=False, italic=False, color=None, caps=False, spacing=None):
    r.font.name = FONT                      # inserts w:rFonts in schema order
    rpr = r._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:cs'), FONT)
    r.font.bold = bold
    r.font.italic = italic
    if caps: r.font.all_caps = True
    if color is not None: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    if spacing is not None:                  # w:spacing must precede w:sz
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing))
        sz = rpr.find(qn('w:sz'))
        if sz is not None: sz.addprevious(sp)
        else: rpr.append(sp)
    return r

def para(space_before=0, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None: p.alignment = align
    return p

def add_text(text, **kw):
    sb = kw.pop('space_before', 0); sa = kw.pop('space_after', 6); al = kw.pop('align', None)
    p = para(sb, sa, al)
    set_run(p.add_run(text), **kw)
    return p

def bottom_rule(p, color=INK_HEX, sz=12):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '6'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    pPr.insert(0, pbdr)   # pBdr precedes spacing/ind/jc in the CT_PPr order

def section_head(num, title):
    p = para(16, 2)
    set_run(p.add_run(num + "   "), size=13, bold=True, color=MIDGREY)
    set_run(p.add_run(title), size=14, bold=True, color=INK)
    bottom_rule(p)

def check(text, tag=None, indent=0.0):
    p = para(0, 3)
    p.paragraph_format.left_indent = Cm(0.6 + indent)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    set_run(p.add_run("☐  "), size=11)
    set_run(p.add_run(text), size=10.5)
    if tag:
        set_run(p.add_run("  " + tag), size=8, bold=True, color=MIDGREY, caps=True)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

_TBLPR_ORDER = ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
                'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
                'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook']

def _tblpr_insert(tblPr, el):
    tag = el.tag.split('}')[1]
    rank = _TBLPR_ORDER.index(tag)
    for child in tblPr:
        ctag = child.tag.split('}')[1]
        if ctag in _TBLPR_ORDER and _TBLPR_ORDER.index(ctag) > rank:
            child.addprevious(el); return
    tblPr.append(el)

def cell_margins(table, top=80, bottom=80, left=120, right=120):
    tblPr = table._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for tag, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); mar.append(e)
    _tblpr_insert(tblPr, mar)

def clean_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    # only thin horizontal insideH + top/bottom; no vertical lines
    for edge in ('top', 'bottom', 'insideH'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'D8D2C4')
        borders.append(e)
    for edge in ('left', 'right', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none'); e.set(qn('w:sz'), '0'); e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'auto')
        borders.append(e)
    _tblpr_insert(tblPr, borders)

def put(cell, text, bold=False, color=INK, size=9.5, italic=False):
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    set_run(cell.paragraphs[0].add_run(text), size=size, bold=bold, color=color, italic=italic)

# ---------------------------------------------------------------- COVER
logo_p = para(0, 4)
logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo_p.add_run().add_picture(LOGO, width=Cm(5.4))

eb = para(10, 2)
set_run(eb.add_run("ANSWER ENGINE OPTIMISATION  ·  CLIENT PLAYBOOK"),
        size=9, bold=True, color=MIDGREY, caps=True, spacing=40)

t = para(0, 2)
set_run(t.add_run("Event & Conference AEO Checklist"), size=26, bold=True, color=INK)

st = para(0, 8)
set_run(st.add_run("How to organise your event for AI visibility — the parts you control by "
                   "code, and the authority you earn off-page."), size=12, italic=True, color=MIDGREY)
bottom_rule(st)

add_text("AI answer engines — ChatGPT, Claude, Perplexity, Google AI Overviews, "
         "Microsoft Copilot and Gemini — increasingly decide what your audience learns "
         "about your event before they ever reach a website. This checklist is the working "
         "method we use to make an organisation the source those engines cite. It separates "
         "the work into what can be solved by code (we deliver it) and the authority that can "
         "only be earned through relationships and time (you own it, we guide it).",
         size=10.5, space_after=8)

# pipeline model
ph = para(6, 2)
set_run(ph.add_run("The citation pipeline"), size=12, bold=True, color=INK)
pl = para(0, 4)
set_run(pl.add_run("CRAWLED   →   INDEXED   →   RETRIEVED for the query   →   "
                   "CHOSEN over rivals   →   CITED"), size=10.5, bold=True, color=INK, spacing=10)
add_text("Every stage must pass. Stages 1–2 are won by code and only get you into the "
         "candidate pool — necessary, but not sufficient. Stages 3–4 are won off-page: "
         "authority, corroboration and the link graph decide whether you are retrieved and "
         "chosen ahead of the official and established sources. Most organisers assume great "
         "content wins. It doesn’t — authoritative, corroborated content wins.",
         size=10.5, space_after=8)

# legend
lg = para(2, 8)
set_run(lg.add_run("Legend:   "), size=9, bold=True, color=INK)
set_run(lg.add_run("[CODE] "), size=9, bold=True, color=INK)
set_run(lg.add_run("controllable on-page — Maxifi delivers.     "), size=9, color=MIDGREY)
set_run(lg.add_run("[OFF] "), size=9, bold=True, color=INK)
set_run(lg.add_run("off-page authority — client-owned, Maxifi advises.     "), size=9, color=MIDGREY)
set_run(lg.add_run("[JOINT] "), size=9, bold=True, color=INK)
set_run(lg.add_run("shared — fastest combined."), size=9, color=MIDGREY)

# ---------------------------------------------------------------- HERO TABLE
section_head("■", "The division of labour — the “bolt-on” model")
add_text("The fastest route to AI visibility combines what only you can grant — authority "
         "— with what we build and run: code, content and measurement.",
         size=10.5, space_after=6)

rows = [
    ("Domain", "You (client) own", "Maxifi delivers"),
    ("Authority  [OFF]", "The official-domain backlink; organising-body endorsement; speaker and partner links",
     "Advisory on who to ask, what anchor text, and which pages to point at"),
    ("Corroboration  [OFF]", "Press relationships; consistency with the official narrative",
     "Content engineered to match and extend the official record so corroboration aligns"),
    ("Link graph  [JOINT]", "Approving links from your owned channels",
     "Distribution scaffolding, anchor strategy and link targets"),
    ("Knowledge base  [JOINT]", "Notability facts and sign-off",
     "Wikidata items and sameAs schema wiring"),
    ("Code / technical  [CODE]", "—",
     "Crawl policy, sitemap, schema, FAQ engineering, entity pages, performance"),
    ("Bolt-on AEO pages  [CODE]", "Source material: programme, speakers, recordings",
     "The canonical, schema-rich, FAQ-first session and speaker records that become the citable source"),
    ("Measurement  [CODE]", "—",
     "Monthly six-engine citation audit, indexation monitoring and reporting"),
]
table = doc.add_table(rows=len(rows), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
clean_borders(table)
cell_margins(table)
widths = (Cm(3.9), Cm(5.6), Cm(7.0))
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        c = table.cell(ri, ci)
        c.width = widths[ci]
        if ri == 0:
            put(c, val, bold=True, color=WHITE, size=10)
            shade(c, INK_HEX)
        else:
            shade(c, CREAM if ri % 2 == 0 else "FFFFFF")
            put(c, val, bold=(ci == 0), color=INK if ci == 0 else RGBColor(0x33,0x33,0x33),
                size=9.5, italic=(val == "—"))

add_text("A “bolt-on” AEO page is a canonical, machine-extractable record for each session "
         "or speaker that sits alongside your existing event site — FAQ-first, schema-complete, "
         "entity-linked, with verified primary detail. You supply the raw material; we build the "
         "page and run the measurement. One official backlink from your domain is what lights it up.",
         size=10, italic=True, color=MIDGREY, space_before=8, space_after=4)

# ---------------------------------------------------------------- PART A
section_head("A", "Parts you control by code  [CODE]")
add_text("On-page and technical AEO. Table stakes — do them all, completely. They get you "
         "crawled and indexed and make you the most extractable source in the candidate pool.",
         size=10, italic=True, color=MIDGREY, space_after=4)

groups_a = [
    ("A1  Crawl access & bot policy", [
        "robots.txt allows all six AI crawlers and variants (GPTBot, OAI-SearchBot, ChatGPT-User; ClaudeBot, Claude-User, Claude-SearchBot; PerplexityBot, Perplexity-User; Googlebot, Google-Extended; Bingbot; Applebot/Applebot-Extended; CCBot).",
        "No accidental Disallow in a named user-agent group (named groups read only their own rules).",
        "Sitemap directive present in robots.txt; no noindex, auth wall, or JS-gated content on citable pages.",
        "llms.txt published at site root listing canonical URLs.",
    ]),
    ("A2  Indexation & discovery", [
        "XML sitemap auto-generated, lists every canonical URL, kept current.",
        "One canonical domain; self-referential canonical tag on every page; no duplicate URLs.",
        "Deep pages submitted via Search Console / Bing URL Inspection to cut indexing latency.",
        "Verified in Google Search Console and Bing Webmaster Tools; shallow, logical URL structure.",
    ]),
    ("A3  Structured data / schema", [
        "Event/BusinessEvent JSON-LD on the event and every session (dates, location, status, organizer, superEvent).",
        "Person schema on every speaker (jobTitle, worksFor, sameAs to LinkedIn/Wikipedia/Wikidata).",
        "FAQPage JSON-LD mirroring visible Q&A verbatim; Organization + WebSite graph site-wide.",
        "BreadcrumbList per page; VideoObject with transcript where recordings exist.",
        "Every block validates (Rich Results Test / schema.org validator) before it is signed off.",
    ]),
    ("A4  Content engineering (extractability)", [
        "Each page answers one clear intent; the lede paragraph is the liftable answer.",
        "FAQ-first — highest-intent questions (who spoke, when/where) answered first in 2–3 sentences.",
        "Self-contained citable facts: each statistic carries a number, a date and an inline source.",
        "Named entities explicit on every mention (full name + role + organisation), never pronouns.",
        "Unique primary information (verified transcripts, quotes, original data), not a rewrite of the official page.",
        "Consistent locale and entity naming; visible publish and last-updated dates.",
    ]),
    ("A5  Entity & internal link architecture", [
        "Permanent canonical entity page for every speaker and session.",
        "Topic clusters: theme/track hubs linking to sessions and back; related-content cross-links.",
        "Stable URLs — never break a canonical record once published.",
    ]),
    ("A6  Performance, rendering & accessibility", [
        "Server-rendered / static HTML (content present without JS execution); fast, mobile-clean, valid markup.",
        "Semantic headings (one H1, ordered H2s); Open Graph + Twitter cards; descriptive image alt text.",
    ]),
    ("A7  On-page provenance & trust", [
        "Explicit authorship/publisher, contact and corrections policy; stated source hierarchy.",
        "Clear independence/affiliation disclosure; a ready-to-use citation string on each record.",
    ]),
]
for title, items in groups_a:
    p = para(8, 2); set_run(p.add_run(title), size=11, bold=True, color=INK)
    for it in items:
        check(it)

# ---------------------------------------------------------------- PART B
section_head("B", "Parts you cannot control by code  [OFF]")
add_text("Off-page authority, corroboration and the link graph — this decides retrieval and "
         "selection (stages 3–4). Earned over time and through relationships, which is exactly "
         "why it is the moat and cannot be faked.",
         size=10, italic=True, color=MIDGREY, space_after=4)

groups_b = [
    ("B1  Domain & entity authority  [OFF]", [
        "Inbound links from authoritative, topical domains (official event site, organising body, regulators, established trade press).",
        "Consistent entity signals across the web (same name, same description); domain age and track record (start early).",
    ]),
    ("B2  Corroboration  [OFF]", [
        "The same facts asserted by multiple independent sources that agree with your page.",
        "Trade-media / news coverage naming the same entities, dates and claims; co-citation alongside trusted pages.",
        "No contradiction with the official record — engines drop conflicting low-authority sources.",
    ]),
    ("B3  Link graph  [OFF] / [JOINT]", [
        "Backlink from the official event domain to the canonical records — highest-leverage single action.  [JOINT]",
        "Backlink from the organising body; speaker/sponsor/partner organisations linking to their session pages.  [JOINT]",
        "Trade-press articles linking to the records; earned links over time as people reference them.  [OFF]",
    ]),
    ("B4  Knowledge-base & entity presence  [JOINT]", [
        "Wikidata item for the event, key speakers and (where notable) sessions; Wikipedia where notability is genuinely met (never fabricate).",
        "Google Knowledge Panel / entity recognition; sameAs links from your schema to those entities (the code half).",
    ]),
    ("B5  Brand, social & distribution  [JOINT]", [
        "Official channels (LinkedIn, YouTube, press wire) publishing and linking back to the records.",
        "Speakers sharing their own session/profile pages; recordings + transcripts hosted and linked; sustained mentions, not a launch burst.",
    ]),
    ("B6  Primary-source & freshness  [OFF] / [JOINT]", [
        "You are the first or fullest publisher of the verified detail (transcript, quotes, data).",
        "Records updated as the official record completes; demonstrable recency cadence so engines re-crawl and re-rank.",
    ]),
]
for title, items in groups_b:
    p = para(8, 2); set_run(p.add_run(title), size=11, bold=True, color=INK)
    for it in items:
        check(it)

# ---------------------------------------------------------------- PART D
section_head("D", "Cadence — when each item happens")
cad = [
    ("8–12 weeks pre-event  [CODE]", "Stand up the hub, schema, sitemap, bot policy and speaker/session scaffolding; verify Search Console + Bing; seed llms.txt."),
    ("Pre-event  [JOINT]", "Secure the official-domain backlink; create Wikidata entities; line up distribution channels."),
    ("During event  [CODE]+[JOINT]", "Publish canonical records live; speakers and partners share their pages; capture recordings."),
    ("0–2 weeks post-event  [CODE]", "Land verified transcripts, quotes and data (become the fullest source); submit deep URLs for indexing."),
    ("Monthly  [CODE]", "Run the six-engine citation audit + indexation check; report the green/grey grid; refresh records on substantive change."),
]
for when, what in cad:
    p = para(3, 2); p.paragraph_format.left_indent = Cm(0.6); p.paragraph_format.first_line_indent = Cm(-0.6)
    set_run(p.add_run("☐  "), size=11)
    set_run(p.add_run(when + " — "), size=10.5, bold=True, color=INK)
    set_run(p.add_run(what), size=10.5)

# ---------------------------------------------------------------- PART E
section_head("E", "Measurement — proving it worked")
for it in [
    "A fixed set of head queries per event (the who, the when/where, the what-was-said questions real users ask).",
    "Each query checked across all six engines monthly; record cited true/false with the verifiable source URL the engine returned.",
    "Never mark a citation without a real source URL — measurement integrity is the product; a fabricated tick destroys the report’s credibility.",
    "Separate “indexed” from “cited” in reporting (different stages, different fixes).",
    "Publish a dated citation report so progress is visible and accountable.",
]:
    check(it)

# ---------------------------------------------------------------- FOOTER
fp = para(18, 0); bottom_rule(fp, color="D8D2C4", sz=6)
f = para(6, 0)
set_run(f.add_run("Maxifi Digital"), size=10, bold=True, color=INK)
set_run(f.add_run("  —  Answer Engine Optimisation.  We build the systems that make expert "
                  "organisations the answer AI gives.   hello@maxifidigital.com  ·  maxifidigital.com"),
        size=9, color=MIDGREY)

doc.save(OUT)
print("Wrote", OUT)
